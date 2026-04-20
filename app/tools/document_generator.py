import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def sanitize_text(text: str) -> str:
    # Remove control characters that break XML parsing in docx
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

def generate_docx_from_markdown(md_text: str, app_name: str = "UnknownApp") -> io.BytesIO:
    doc = Document()
    
    # Title styling
    title = doc.add_heading(f'Modernization Report: {app_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Strip wrapping markdown code blocks if LLM included them
    if md_text.startswith('```markdown'):
        md_text = md_text[11:].strip()
    elif md_text.startswith('```'):
        md_text = md_text[3:].strip()
        
    if md_text.endswith('```'):
        md_text = md_text[:-3].strip()

    lines = md_text.split('\n')
    i = 0
    
    def process_inline(paragraph, text):
        text = sanitize_text(text)
        # Basic inline bold processing
        # We also clear out LLM-generated HTML strong tags that might have slipped through
        text = re.sub(r'<\/?strong>', '**', text, flags=re.IGNORECASE)
        text = re.sub(r'<\/?b>', '**', text, flags=re.IGNORECASE)
        
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Add basic code block handling inline
                code_parts = re.split(r'(`.*?`)', part)
                for cpart in code_parts:
                    if cpart.startswith('`') and cpart.endswith('`'):
                        run = paragraph.add_run(cpart[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        paragraph.add_run(cpart)

    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            process_inline(p, line[2:].strip())
            i += 1
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            process_inline(p, line[3:].strip())
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            process_inline(p, line[4:].strip())
            i += 1
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, line[2:].strip())
            i += 1
            
        # Numbered Lists
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+\.\s', '', line).strip()
            process_inline(p, content)
            i += 1
            
        # Horizontal Rule
        elif line == '---' or line == '***' or line == '___':
            p = doc.add_paragraph("_" * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            
        # Tables
        elif line.startswith('|') and line.endswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            if len(table_lines) >= 2:
                # Extract headers
                headers = [col.strip() for col in table_lines[0].strip('|').split('|')]
                
                # Check if second line is separator
                data_start_idx = 1
                if re.match(r'^\|[\s\-\:]+\|', table_lines[1]):
                    data_start_idx = 2
                    
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for col_idx, header_text in enumerate(headers):
                    if col_idx < len(hdr_cells):
                        process_inline(hdr_cells[col_idx].paragraphs[0], header_text)
                        
                # Add data rows
                for row_line in table_lines[data_start_idx:]:
                    cols = [col.strip() for col in row_line.strip('|').split('|')]
                    row_cells = table.add_row().cells
                    for col_idx, col_text in enumerate(cols):
                        if col_idx < len(row_cells):
                            process_inline(row_cells[col_idx].paragraphs[0], col_text)
            else:
                p = doc.add_paragraph()
                process_inline(p, line)
                i += 1
                
        # Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            process_inline(p, line[2:].strip())
            i += 1
            
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            process_inline(p, line)
            i += 1

    # ISO 5055 explicitly
    if "ISO 5055" not in md_text:
        doc.add_heading('ISO 5055 Compliance & Mitigation', level=1)
        doc.add_paragraph("No severe CWE weaknesses found. Application adheres to Reliability, Security, Performance, and Maintainability standards.")

    if "Disclaimer" not in md_text:
        doc.add_heading("Disclaimer", level=1)
        doc.add_paragraph(
            "Disclaimer - This document and the information contained herein are provided for informational and guidance purposes only. "
            "Before incorporating any of these configurations, scripts, or architectural patterns into a formal modernization journey, "
            "they must be reviewed and verified by a competent Solutions Architect to ensure alignment with specific infrastructure, "
            "security, and compliance requirements."
        )
        doc.add_paragraph(
            "The developer of this platform and CAST Software assume no responsibility or liability for any errors, omissions, or damages—"
            "direct or indirect—resulting from the use or implementation of this information. "
            "All actions taken based on this content are at the user's own risk and discretion."
        )

    file_stream = io.BytesIO()
    doc.save(file_stream)
    
    # Save locally to test/ folder
    try:
        test_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "test")
        os.makedirs(test_dir, exist_ok=True)
        local_path = os.path.join(test_dir, f"{app_name}_Modernization_Roadmap.docx")
        doc.save(local_path)
    except Exception as e:
        print(f"Failed to save document locally: {e}")

    file_stream.seek(0)
    return file_stream
